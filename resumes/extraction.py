import logging
import re
import unicodedata
from typing import Optional

logger = logging.getLogger(__name__)


class ExtractionError(Exception):
    """Base exception for text extraction errors."""
    def __init__(self, message: str, error_code: str):
        super().__init__(message)
        self.error_code = error_code


def _extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using pdfminer.
    
    Raises:
        ExtractionError: With specific error codes for different failure modes.
    """
    try:
        from pdfminer.high_level import extract_text
        from pdfminer.pdfparser import PDFSyntaxError
        from pdfminer.pdfdocument import PDFEncryptionError
    except ImportError as e:
        logger.error("pdfminer not installed", extra={"error": str(e)})
        raise ExtractionError("PDF extraction library not available", "MISSING_DEPENDENCY")
    
    try:
        text = extract_text(file_path) or ""
        if not text.strip():
            logger.warning("PDF extraction returned empty text", extra={"file_path": file_path})
        return text
    except PDFEncryptionError:
        logger.warning("PDF is password protected", extra={"file_path": file_path})
        raise ExtractionError("PDF is password protected", "PASSWORD_PROTECTED")
    except PDFSyntaxError as e:
        logger.warning("PDF syntax error (corrupted)", extra={"file_path": file_path, "error": str(e)})
        raise ExtractionError(f"PDF appears to be corrupted: {str(e)}", "CORRUPTED_PDF")
    except Exception as e:
        logger.error("PDF extraction failed", extra={"file_path": file_path, "error": str(e)})
        raise ExtractionError(f"PDF extraction failed: {str(e)}", "PDF_EXTRACTION_ERROR")


def _extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx.
    
    Raises:
        ExtractionError: With specific error codes for different failure modes.
    """
    try:
        import docx
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as e:
        logger.error("python-docx not installed", extra={"error": str(e)})
        raise ExtractionError("DOCX extraction library not available", "MISSING_DEPENDENCY")
    
    try:
        doc = docx.Document(file_path)
        parts = []
        
        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
                if row_text:
                    parts.append(row_text)
        
        # Extract headers and footers
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    if p.text and p.text not in parts:
                        parts.insert(0, p.text)
        
        text = "\n".join(parts)
        if not text.strip():
            logger.warning("DOCX extraction returned empty text", extra={"file_path": file_path})
        return text
        
    except PackageNotFoundError:
        logger.warning("DOCX file not found or invalid", extra={"file_path": file_path})
        raise ExtractionError("DOCX file is invalid or corrupted", "CORRUPTED_DOCX")
    except Exception as e:
        logger.error("DOCX extraction failed", extra={"file_path": file_path, "error": str(e)})
        raise ExtractionError(f"DOCX extraction failed: {str(e)}", "DOCX_EXTRACTION_ERROR")


def _extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from legacy .doc file using docx2txt.
    
    Raises:
        ExtractionError: With specific error codes for different failure modes.
    """
    try:
        import docx2txt
    except ImportError as e:
        logger.error("docx2txt not installed", extra={"error": str(e)})
        raise ExtractionError("Legacy DOC extraction library not available", "MISSING_DEPENDENCY")
    
    try:
        text = docx2txt.process(file_path) or ""
        if not text.strip():
            logger.warning("DOC extraction returned empty text", extra={"file_path": file_path})
        return text
    except Exception as e:
        logger.error("DOC extraction failed", extra={"file_path": file_path, "error": str(e)})
        raise ExtractionError(f"DOC extraction failed: {str(e)}", "DOC_EXTRACTION_ERROR")


def _extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from plain text file with encoding detection.
    
    Raises:
        ExtractionError: With specific error codes for different failure modes.
    """
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error("TXT file read error", extra={"file_path": file_path, "error": str(e)})
            raise ExtractionError(f"Failed to read text file: {str(e)}", "TXT_READ_ERROR")
    
    # Fallback: read with replacement
    logger.warning("TXT encoding detection failed, using fallback", extra={"file_path": file_path})
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text for better LLM processing.
    
    Improvements:
    - Unicode normalization (NFKD)
    - Smart quote replacement
    - Control character removal
    - Hyphenated word fixing
    - Whitespace normalization
    """
    if not text:
        return ""
    
    # Unicode normalization (decompose characters)
    text = unicodedata.normalize('NFKC', text)
    
    # Replace smart quotes with standard quotes
    smart_quotes = {
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote
        '\u2013': '-',  # En dash
        '\u2014': '-',  # Em dash
        '\u2026': '...',  # Ellipsis
        '\u00a0': ' ',  # Non-breaking space
        '\u200b': '',   # Zero-width space
        '\u200c': '',   # Zero-width non-joiner
        '\u200d': '',   # Zero-width joiner
        '\ufeff': '',   # BOM
    }
    for char, replacement in smart_quotes.items():
        text = text.replace(char, replacement)
    
    # Remove null bytes and other control characters (except newlines and tabs)
    text = ''.join(c if c in '\n\t' or (ord(c) >= 32 and ord(c) != 127) else ' ' for c in text)
    
    # Fix hyphenated words split across lines (common in PDFs)
    text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
    
    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
    text = re.sub(r' *\n *', '\n', text)  # Remove spaces around newlines
    text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines
    
    # Strip leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text.strip()


def extract_text_from_file(file_path: str, mime_type: str, original_filename: str) -> tuple[str, str]:
    """
    Dispatches to the correct extraction method based on file extension/mime type.
    Returns (raw_text, extraction_method).
    """
    name = (original_filename or "").lower()
    
    if name.endswith(".pdf") or mime_type == "application/pdf":
        return _extract_text_from_pdf(file_path), "pdfminer"
    
    if name.endswith(".docx") or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_text_from_docx(file_path), "python-docx"
    
    if name.endswith(".doc") or mime_type == "application/msword":
        return _extract_text_from_doc(file_path), "docx2txt"
    
    if name.endswith(".txt") or mime_type == "text/plain":
        return _extract_text_from_txt(file_path), "plaintext"
    
    # Default fallback
    logger.warning("Unknown file type, attempting DOCX extraction", extra={
        "file_name": name,
    })
    return _extract_text_from_docx(file_path), "python-docx"
