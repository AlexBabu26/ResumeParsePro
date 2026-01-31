import hashlib
import logging
import re
import unicodedata

from django.conf import settings
from rest_framework import permissions, status, viewsets
from rest_framework.decorators import action
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework.throttling import ScopedRateThrottle

from core.responses import ok, fail

from .models import ResumeDocument, ParseRun
from .serializers import ResumeDocumentSerializer, ResumeUploadSerializer, BulkResumeUploadSerializer, ParseRunSerializer
from .tasks import parse_resume_parse_run

from candidates.models import Candidate

logger = logging.getLogger(__name__)


class ExtractionError(Exception):
    """Base exception for text extraction errors."""
    def __init__(self, message: str, error_code: str):
        super().__init__(message)
        self.error_code = error_code


def _extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using pdfminer.
    
    Raises:
        ExtractionError: With specific error codes for different failure modes.
    """
    try:
        from pdfminer.high_level import extract_text
        from pdfminer.pdfparser import PDFSyntaxError
        from pdfminer.pdfdocument import PDFEncryptionError
    except ImportError as e:
        logger.error("pdfminer not installed", extra={"error": str(e)})
        raise ExtractionError("PDF extraction library not available", "MISSING_DEPENDENCY")
    
    try:
        text = extract_text(file_path) or ""
        if not text.strip():
            logger.warning("PDF extraction returned empty text", extra={"file_path": file_path})
        return text
    except PDFEncryptionError:
        logger.warning("PDF is password protected", extra={"file_path": file_path})
        raise ExtractionError("PDF is password protected", "PASSWORD_PROTECTED")
    except PDFSyntaxError as e:
        logger.warning("PDF syntax error (corrupted)", extra={"file_path": file_path, "error": str(e)})
        raise ExtractionError(f"PDF appears to be corrupted: {str(e)}", "CORRUPTED_PDF")
    except Exception as e:
        logger.error("PDF extraction failed", extra={"file_path": file_path, "error": str(e)})
        raise ExtractionError(f"PDF extraction failed: {str(e)}", "PDF_EXTRACTION_ERROR")


def _extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx.
    
    Raises:
        ExtractionError: With specific error codes for different failure modes.
    """
    try:
        import docx
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as e:
        logger.error("python-docx not installed", extra={"error": str(e)})
        raise ExtractionError("DOCX extraction library not available", "MISSING_DEPENDENCY")
    
    try:
        doc = docx.Document(file_path)
        parts = []
        
        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
                if row_text:
                    parts.append(row_text)
        
        # Extract headers and footers
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    if p.text and p.text not in parts:
                        parts.insert(0, p.text)
        
        text = "\n".join(parts)
        if not text.strip():
            logger.warning("DOCX extraction returned empty text", extra={"file_path": file_path})
        return text
        
    except PackageNotFoundError:
        logger.warning("DOCX file not found or invalid", extra={"file_path": file_path})
        raise ExtractionError("DOCX file is invalid or corrupted", "CORRUPTED_DOCX")
    except Exception as e:
        logger.error("DOCX extraction failed", extra={"file_path": file_path, "error": str(e)})
        raise ExtractionError(f"DOCX extraction failed: {str(e)}", "DOCX_EXTRACTION_ERROR")


def _extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from legacy .doc file using docx2txt.
    
    Raises:
        ExtractionError: With specific error codes for different failure modes.
    """
    try:
        import docx2txt
    except ImportError as e:
        logger.error("docx2txt not installed", extra={"error": str(e)})
        raise ExtractionError("Legacy DOC extraction library not available", "MISSING_DEPENDENCY")
    
    try:
        text = docx2txt.process(file_path) or ""
        if not text.strip():
            logger.warning("DOC extraction returned empty text", extra={"file_path": file_path})
        return text
    except Exception as e:
        logger.error("DOC extraction failed", extra={"file_path": file_path, "error": str(e)})
        raise ExtractionError(f"DOC extraction failed: {str(e)}", "DOC_EXTRACTION_ERROR")


def _extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from plain text file with encoding detection.
    
    Raises:
        ExtractionError: With specific error codes for different failure modes.
    """
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error("TXT file read error", extra={"file_path": file_path, "error": str(e)})
            raise ExtractionError(f"Failed to read text file: {str(e)}", "TXT_READ_ERROR")
    
    # Fallback: read with replacement
    logger.warning("TXT encoding detection failed, using fallback", extra={"file_path": file_path})
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def _clean_text(text: str) -> str:
    """
    Clean and normalize extracted text for better LLM processing.
    
    Improvements:
    - Unicode normalization (NFKD)
    - Smart quote replacement
    - Control character removal
    - Hyphenated word fixing
    - Whitespace normalization
    """
    if not text:
        return ""
    
    # Unicode normalization (decompose characters)
    text = unicodedata.normalize('NFKC', text)
    
    # Replace smart quotes with standard quotes
    smart_quotes = {
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote
        '\u2013': '-',  # En dash
        '\u2014': '-',  # Em dash
        '\u2026': '...',  # Ellipsis
        '\u00a0': ' ',  # Non-breaking space
        '\u200b': '',   # Zero-width space
        '\u200c': '',   # Zero-width non-joiner
        '\u200d': '',   # Zero-width joiner
        '\ufeff': '',   # BOM
    }
    for char, replacement in smart_quotes.items():
        text = text.replace(char, replacement)
    
    # Remove null bytes and other control characters (except newlines and tabs)
    text = ''.join(c if c in '\n\t' or (ord(c) >= 32 and ord(c) != 127) else ' ' for c in text)
    
    # Fix hyphenated words split across lines (common in PDFs)
    text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
    
    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
    text = re.sub(r' *\n *', '\n', text)  # Remove spaces around newlines
    text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines
    
    # Strip leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text.strip()


def sha256_of_uploaded_file(uploaded_file) -> str:
    h = hashlib.sha256()
    pos = uploaded_file.tell() if hasattr(uploaded_file, "tell") else None
    for chunk in uploaded_file.chunks():
        h.update(chunk)
    # reset pointer so Django can save it
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(pos or 0)
    return h.hexdigest()


def _calculate_years_experience(candidate: Candidate) -> float:
    """Calculate total years of experience from experience entries"""
    from datetime import datetime
    
    try:
        from dateutil import parser
        use_dateutil = True
    except ImportError:
        use_dateutil = False
    
    total_years = 0.0
    for exp in candidate.experience.all():
        if not exp.start_date:
            continue
        
        try:
            if use_dateutil:
                start = parser.parse(exp.start_date, default=datetime(2000, 1, 1))
                end = parser.parse(exp.end_date, default=datetime.now()) if exp.end_date else datetime.now()
            else:
                # Fallback: simple date parsing (YYYY-MM-DD format)
                start_str = str(exp.start_date)[:10]  # Take first 10 chars (YYYY-MM-DD)
                end_str = str(exp.end_date)[:10] if exp.end_date else None
                
                try:
                    start = datetime.strptime(start_str, "%Y-%m-%d")
                    end = datetime.strptime(end_str, "%Y-%m-%d") if end_str else datetime.now()
                except ValueError:
                    # If parsing fails, skip this entry
                    continue
            
            # Calculate years
            delta = end - start
            years = delta.days / 365.25
            total_years += max(0, years)
        except (ValueError, TypeError):
            # If date parsing fails, skip this entry
            continue
    
    return total_years


def _build_candidate_data_for_validation(candidate: Candidate) -> dict:
    """Build a dictionary of candidate data for LLM validation."""
    return {
        "full_name": candidate.full_name,
        "primary_role": candidate.primary_role,
        "seniority": candidate.seniority,
        "location": candidate.location,
        "headline": candidate.headline,
        "skills": [{"name": s.name, "category": s.category} for s in candidate.skills.all()],
        "education": [
            {
                "institution": ed.institution,
                "degree": ed.degree,
                "field_of_study": ed.field_of_study,
                "start_date": ed.start_date,
                "end_date": ed.end_date,
            }
            for ed in candidate.education.all()
        ],
        "experience": [
            {
                "company": ex.company,
                "title": ex.title,
                "start_date": ex.start_date,
                "end_date": ex.end_date,
                "is_current": ex.is_current,
            }
            for ex in candidate.experience.all()
        ],
        "overall_confidence": candidate.overall_confidence,
        "summary_one_liner": candidate.summary_one_liner,
    }


def _candidate_meets_requirements_llm(candidate: Candidate, requirements: dict) -> tuple[bool, list[str]]:
    """
    Use LLM to check if a candidate meets the specified requirements.
    More accurate than string matching, especially for role comparisons.
    Returns (meets_requirements: bool, reasons: list[str])
    """
    from .pipeline import call_requirements_validation
    
    candidate_data = _build_candidate_data_for_validation(candidate)
    
    try:
        result = call_requirements_validation(candidate_data, requirements)
        meets = result.get("meets_requirements", False)
        reasons = result.get("reasons", [])
        if not isinstance(reasons, list):
            reasons = [str(reasons)] if reasons else []
        return meets, reasons
    except Exception as e:
        # Fallback to string-based validation if LLM fails
        return _candidate_meets_requirements_string(candidate, requirements)


def _candidate_meets_requirements_string(candidate: Candidate, requirements: dict) -> tuple[bool, list[str]]:
    """
    String-based validation (fallback method).
    Less accurate but faster and doesn't require LLM call.
    """
    reasons = []
    meets = True
    
    # Check required skills (all must be present)
    if "required_skills" in requirements:
        required = [s.lower() for s in requirements["required_skills"]]
        candidate_skills = [s.name.lower() for s in candidate.skills.all()]
        missing = [s for s in required if s not in candidate_skills]
        if missing:
            meets = False
            reasons.append(f"Missing required skills: {', '.join(missing)}")
    
    # Check any skills (at least one must be present)
    if "any_skills" in requirements:
        any_skills = [s.lower() for s in requirements["any_skills"]]
        candidate_skills = [s.name.lower() for s in candidate.skills.all()]
        if not any(s in candidate_skills for s in any_skills):
            meets = False
            reasons.append(f"Missing at least one of these skills: {', '.join(requirements['any_skills'])}")
    
    # Check minimum years of experience
    if "min_years_experience" in requirements:
        years = _calculate_years_experience(candidate)
        if years < requirements["min_years_experience"]:
            meets = False
            reasons.append(f"Insufficient experience: {years:.1f} years (required: {requirements['min_years_experience']})")
    
    # Check required education degree
    if "required_education_degree" in requirements:
        required_degrees = [d.lower() for d in requirements["required_education_degree"]]
        candidate_degrees = [ed.degree.lower() if ed.degree else "" for ed in candidate.education.all()]
        if not any(degree and any(rd in degree.lower() for rd in required_degrees) for degree in candidate_degrees):
            meets = False
            reasons.append(f"Missing required education degree: {', '.join(requirements['required_education_degree'])}")
    
    # Check required primary role
    if "required_primary_role" in requirements:
        required_roles = [r.lower().strip() for r in requirements["required_primary_role"]]
        candidate_role = (candidate.primary_role or "").lower().strip()
        
        if not candidate_role:
            meets = False
            reasons.append(f"Primary role not found (required: {', '.join(requirements['required_primary_role'])})")
        else:
            # Strict matching: check if any required role is a substring of candidate role
            # OR if candidate role is a substring of any required role
            role_matches = any(
                required_role in candidate_role or candidate_role in required_role
                for required_role in required_roles
            )
            
            if not role_matches:
                meets = False
                reasons.append(f"Primary role mismatch: '{candidate.primary_role}' (required: {', '.join(requirements['required_primary_role'])})")
    
    # Check required seniority
    if "required_seniority" in requirements:
        required_seniorities = [s.lower() for s in requirements["required_seniority"]]
        candidate_seniority = (candidate.seniority or "").lower()
        if candidate_seniority not in required_seniorities:
            meets = False
            reasons.append(f"Seniority mismatch: '{candidate.seniority}' (required: {', '.join(requirements['required_seniority'])})")
    
    # Check location contains
    if "location_contains" in requirements:
        location_search = requirements["location_contains"].lower()
        candidate_location = (candidate.location or "").lower()
        if location_search not in candidate_location:
            meets = False
            reasons.append(f"Location mismatch: '{candidate.location}' (must contain: '{requirements['location_contains']}')")
    
    # Check minimum confidence
    if "min_confidence" in requirements:
        if candidate.overall_confidence < requirements["min_confidence"]:
            meets = False
            reasons.append(f"Low confidence: {candidate.overall_confidence:.2f} (required: {requirements['min_confidence']})")
    
    return meets, reasons


def _candidate_meets_requirements(candidate: Candidate, requirements: dict, use_llm: bool = True) -> tuple[bool, list[str]]:
    """
    Check if a candidate meets the specified requirements.
    Returns (meets_requirements: bool, reasons: list[str])
    
    Args:
        candidate: The candidate to check
        requirements: Dictionary of requirements to check against
        use_llm: If True, use LLM for semantic validation (more accurate but slower)
                 If False, use string-based validation (faster but less accurate)
    """
    if not requirements:
        return True, []
    
    # Check if LLM validation is requested (default: True for accuracy)
    if use_llm and requirements.get("use_llm_validation", True):
        return _candidate_meets_requirements_llm(candidate, requirements)
    else:
        return _candidate_meets_requirements_string(candidate, requirements)


class ResumeDocumentViewSet(viewsets.ModelViewSet):
    serializer_class = ResumeDocumentSerializer
    permission_classes = [permissions.IsAuthenticated]
    # Only allow list, retrieve, destroy (no create/update via this viewset)
    http_method_names = ['get', 'delete', 'head', 'options']

    def get_queryset(self):
        # (3) Ownership filtering
        return ResumeDocument.objects.filter(uploaded_by=self.request.user).order_by("-created_at")

    def destroy(self, request, *args, **kwargs):
        """Delete a resume document and all associated parse runs and candidates (cascade)."""
        instance = self.get_object()
        doc_id = instance.id
        filename = instance.original_filename
        
        # Count related objects before deletion (for logging)
        parse_runs_count = instance.parse_runs.count()
        candidates_count = instance.candidate_profiles.count()
        
        # Delete the file from filesystem if it exists
        if instance.file:
            try:
                instance.file.delete(save=False)
            except Exception as e:
                logger.warning(f"Failed to delete file for ResumeDocument {doc_id}: {e}")
        
        # Delete the document (cascade will delete parse_runs and candidates)
        instance.delete()
        
        logger.info(f"ResumeDocument {doc_id} deleted by user {request.user.id}", extra={
            "document_id": doc_id,
            "file_name": filename,
            "parse_runs_deleted": parse_runs_count,
            "candidates_deleted": candidates_count,
        })
        
        return ok({
            "message": f"Document '{filename}' deleted successfully",
            "parse_runs_deleted": parse_runs_count,
            "candidates_deleted": candidates_count,
        })


class ParseRunViewSet(viewsets.ModelViewSet):
    serializer_class = ParseRunSerializer
    permission_classes = [permissions.IsAuthenticated]
    throttle_classes = [ScopedRateThrottle]
    throttle_scope = "parse_retry"
    # Only allow list, retrieve, destroy, and retry action (no create/update via this viewset)
    http_method_names = ['get', 'post', 'delete', 'head', 'options']

    def get_queryset(self):
        # (3) Ownership filtering
        return ParseRun.objects.select_related("resume_document").filter(
            resume_document__uploaded_by=self.request.user
        ).order_by("-created_at")

    def destroy(self, request, *args, **kwargs):
        """Delete a parse run and its associated candidate (if any)."""
        instance = self.get_object()
        run_id = instance.id
        
        # Also delete associated candidate if exists
        from candidates.models import Candidate
        deleted_candidate = Candidate.objects.filter(parse_run=instance).delete()[0]
        
        instance.delete()
        logger.info(f"ParseRun {run_id} deleted by user {request.user.id}")
        
        return ok({
            "message": f"Parse run #{run_id} deleted successfully",
            "deleted_candidate": deleted_candidate > 0
        })

    @action(detail=True, methods=["post"])
    def retry(self, request, pk=None):
        run = self.get_object()
        doc = run.resume_document

        if not doc.raw_text:
            return fail("No raw_text available for retry.", code="NO_RAW_TEXT", status=400)

        new_run = ParseRun.objects.create(
            resume_document=doc,
            status="queued",
            model_name=getattr(settings, "OPENROUTER_EXTRACT_MODEL", "openai/gpt-4o-mini"),
            prompt_version="v1",
            temperature=float(getattr(settings, "OPENROUTER_TEMPERATURE", 0.1)),
        )

        # dispatch async, or sync if requested
        sync = request.query_params.get("sync") == "1"
        if getattr(settings, "RESUME_PARSE_ASYNC", True) and not sync:
            parse_resume_parse_run.delay(new_run.id)
            return ok({"parse_run_id": new_run.id, "status": "queued"}, status=202)

        # synchronous fallback
        parse_resume_parse_run(new_run.id)
        new_run.refresh_from_db()
        return ok(ParseRunSerializer(new_run).data, status=201)


class ResumeUploadViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]
    throttle_classes = [ScopedRateThrottle]
    throttle_scope = "resumes_upload"

    @action(detail=False, methods=["post"], url_path="upload")
    def upload(self, request):
        ser = ResumeUploadSerializer(data=request.data)
        ser.is_valid(raise_exception=True)
        f = ser.validated_data["file"]

        # Get requirements from request data (JSON field) - support for single file upload
        requirements_json = request.data.get("requirements")
        if requirements_json and isinstance(requirements_json, str):
            import json
            try:
                requirements_json = json.loads(requirements_json)
            except json.JSONDecodeError:
                return fail("Invalid JSON in requirements field.", code="INVALID_REQUIREMENTS", status=400)
        
        # Validate requirements if provided
        requirements = None
        if requirements_json:
            from .serializers import BulkResumeUploadSerializer
            temp_ser = BulkResumeUploadSerializer(data={"files": [f], "requirements": requirements_json})
            if temp_ser.is_valid():
                requirements = temp_ser.validated_data.get("requirements")
            else:
                return fail(f"Invalid requirements: {temp_ser.errors}", code="INVALID_REQUIREMENTS", status=400)

        # (2) Idempotency: compute hash and check duplicates per user
        file_hash = sha256_of_uploaded_file(f)
        existing = ResumeDocument.objects.filter(uploaded_by=request.user, file_hash=file_hash).first()
        if existing:
            latest_candidate = Candidate.objects.filter(resume_document=existing).order_by("-created_at").first()
            latest_run = existing.parse_runs.order_by("-created_at").first()
            
            # Check requirements for duplicates too (sync mode only)
            if requirements and latest_candidate:
                sync = request.query_params.get("sync") == "1"
                if sync:
                    meets, reasons = _candidate_meets_requirements(latest_candidate, requirements)
                    if not meets:
                        return ok({
                            "duplicate": True,
                            "resume_document_id": existing.id,
                            "parse_run_id": latest_run.id if latest_run else None,
                            "candidate_id": latest_candidate.id if latest_candidate else None,
                            "status": latest_run.status if latest_run else None,
                            "requirements_check": "failed",
                            "rejection_reasons": reasons,
                        }, status=200)
            
            return ok({
                "duplicate": True,
                "resume_document_id": existing.id,
                "parse_run_id": latest_run.id if latest_run else None,
                "candidate_id": latest_candidate.id if latest_candidate else None,
                "status": latest_run.status if latest_run else None,
            }, status=200)

        # Save document
        doc = ResumeDocument.objects.create(
            original_filename=f.name,
            file=f,
            mime_type=getattr(f, "content_type", "") or "",
            file_hash=file_hash,
            file_size=getattr(f, "size", 0) or 0,
            uploaded_by=request.user,
        )

        # Extract text now (fast and deterministic)
        try:
            file_path = doc.file.path
            name = (doc.original_filename or "").lower()
            
            logger.info("Starting text extraction", extra={
                "document_id": doc.id,
                "file_name": doc.original_filename,
                "mime_type": doc.mime_type,
            })

            if name.endswith(".pdf") or doc.mime_type == "application/pdf":
                raw = _extract_text_from_pdf(file_path)
                doc.extraction_method = "pdfminer"
            elif name.endswith(".docx") or doc.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                raw = _extract_text_from_docx(file_path)
                doc.extraction_method = "python-docx"
            elif name.endswith(".doc") or doc.mime_type == "application/msword":
                raw = _extract_text_from_doc(file_path)
                doc.extraction_method = "docx2txt"
            elif name.endswith(".txt") or doc.mime_type == "text/plain":
                raw = _extract_text_from_txt(file_path)
                doc.extraction_method = "plaintext"
            else:
                # Default to DOCX extraction for unknown types
                logger.warning("Unknown file type, attempting DOCX extraction", extra={
                    "document_id": doc.id,
                    "file_name": name,
                })
                raw = _extract_text_from_docx(file_path)
                doc.extraction_method = "python-docx"

            doc.raw_text = _clean_text(raw)
            doc.save(update_fields=["raw_text", "extraction_method", "updated_at"])
            
            logger.info("Text extraction complete", extra={
                "document_id": doc.id,
                "extraction_method": doc.extraction_method,
                "text_length": len(doc.raw_text),
            })
            
        except ExtractionError as e:
            doc.raw_text = ""
            doc.extraction_method = "failed"
            doc.save(update_fields=["raw_text", "extraction_method", "updated_at"])
            logger.warning("Text extraction failed", extra={
                "document_id": doc.id,
                "error_code": e.error_code,
                "error": str(e),
            })
            return fail(str(e), code=e.error_code, status=400)
        except Exception as e:
            doc.raw_text = ""
            doc.extraction_method = "failed"
            doc.save(update_fields=["raw_text", "extraction_method", "updated_at"])
            logger.error("Unexpected extraction error", extra={
                "document_id": doc.id,
                "error": str(e),
            })
            return fail(f"Text extraction failed: {str(e)}", code="TEXT_EXTRACTION_FAILED", status=400)

        # Create ParseRun queued and dispatch (7)
        run = ParseRun.objects.create(
            resume_document=doc,
            status="queued",
            model_name=getattr(settings, "OPENROUTER_EXTRACT_MODEL", "openai/gpt-4o-mini"),
            prompt_version="v1",
            temperature=float(getattr(settings, "OPENROUTER_TEMPERATURE", 0.1)),
        )

        sync = request.query_params.get("sync") == "1"
        if getattr(settings, "RESUME_PARSE_ASYNC", True) and not sync:
            parse_resume_parse_run.delay(run.id)
            return ok({"resume_document_id": doc.id, "parse_run_id": run.id, "status": "queued"}, status=202)

        # sync fallback for demos/testing
        parse_resume_parse_run(run.id)
        run.refresh_from_db()
        latest_candidate = Candidate.objects.filter(parse_run=run).order_by("-created_at").first()
        
        # Check requirements if provided (sync mode only)
        rejected = False
        rejection_reasons = []
        if requirements and latest_candidate:
            meets, reasons = _candidate_meets_requirements(latest_candidate, requirements)
            if not meets:
                rejected = True
                rejection_reasons = reasons
                latest_candidate.delete()  # Discard candidate that doesn't meet requirements
                latest_candidate = None
        
        response_data = {
            "resume_document_id": doc.id,
            "parse_run_id": run.id,
            "status": run.status,
            "candidate_id": latest_candidate.id if latest_candidate else None,
        }
        
        if requirements:
            response_data["requirements_applied"] = requirements
            if rejected:
                response_data["rejected"] = True
                response_data["rejection_reasons"] = rejection_reasons
            else:
                response_data["accepted"] = True
        
        return ok(response_data, status=201)

    @action(detail=False, methods=["post"], url_path="bulk-upload")
    def bulk_upload(self, request):
        """
        Bulk upload endpoint that accepts multiple files.
        Returns a summary of all uploads with their status.
        """
        # Handle both 'files' (list) and 'file' (multiple files with same name)
        files = request.FILES.getlist("files") or request.FILES.getlist("file")
        
        if not files:
            return fail("No files provided. Use 'files' field for multiple files.", code="NO_FILES", status=400)
        
        if len(files) > 100:
            return fail("Maximum 100 files allowed per bulk upload.", code="TOO_MANY_FILES", status=400)

        # Get requirements from request data (JSON field)
        requirements_json = request.data.get("requirements")
        if requirements_json and isinstance(requirements_json, str):
            import json
            try:
                requirements_json = json.loads(requirements_json)
            except json.JSONDecodeError:
                return fail("Invalid JSON in requirements field.", code="INVALID_REQUIREMENTS", status=400)

        ser = BulkResumeUploadSerializer(data={"files": files, "requirements": requirements_json})
        ser.is_valid(raise_exception=True)
        validated_files = ser.validated_data["files"]
        requirements = ser.validated_data.get("requirements")

        sync = request.query_params.get("sync") == "1"
        results = []
        errors = []
        discarded = []  # Candidates that don't meet requirements

        for idx, f in enumerate(validated_files):
            try:
                # Idempotency: compute hash and check duplicates per user
                file_hash = sha256_of_uploaded_file(f)
                existing = ResumeDocument.objects.filter(uploaded_by=request.user, file_hash=file_hash).first()
                
                if existing:
                    latest_candidate = Candidate.objects.filter(resume_document=existing).order_by("-created_at").first()
                    latest_run = existing.parse_runs.order_by("-created_at").first()
                    
                    # Check requirements for duplicates too (sync mode only)
                    if requirements and latest_candidate and sync:
                        meets, reasons = _candidate_meets_requirements(latest_candidate, requirements)
                        if not meets:
                            discarded.append({
                                "filename": f.name,
                                "candidate_id": latest_candidate.id,
                                "reasons": reasons,
                                "duplicate": True,
                            })
                            results.append({
                                "filename": f.name,
                                "duplicate": True,
                                "resume_document_id": existing.id,
                                "parse_run_id": latest_run.id if latest_run else None,
                                "candidate_id": latest_candidate.id if latest_candidate else None,
                                "status": latest_run.status if latest_run else None,
                                "discarded": True,
                                "discard_reasons": reasons,
                            })
                            continue
                    
                    # If duplicate meets requirements or no requirements, add to results
                    results.append({
                        "filename": f.name,
                        "duplicate": True,
                        "resume_document_id": existing.id,
                        "parse_run_id": latest_run.id if latest_run else None,
                        "candidate_id": latest_candidate.id if latest_candidate else None,
                        "status": latest_run.status if latest_run else None,
                    })
                    continue

                # Save document
                doc = ResumeDocument.objects.create(
                    original_filename=f.name,
                    file=f,
                    mime_type=getattr(f, "content_type", "") or "",
                    file_hash=file_hash,
                    file_size=getattr(f, "size", 0) or 0,
                    uploaded_by=request.user,
                )

                # Extract text now (fast and deterministic)
                try:
                    file_path = doc.file.path
                    name = (doc.original_filename or "").lower()

                    if name.endswith(".pdf") or doc.mime_type == "application/pdf":
                        raw = _extract_text_from_pdf(file_path)
                        doc.extraction_method = "pdfminer"
                    elif name.endswith(".docx") or doc.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        raw = _extract_text_from_docx(file_path)
                        doc.extraction_method = "python-docx"
                    elif name.endswith(".doc") or doc.mime_type == "application/msword":
                        raw = _extract_text_from_doc(file_path)
                        doc.extraction_method = "docx2txt"
                    elif name.endswith(".txt") or doc.mime_type == "text/plain":
                        raw = _extract_text_from_txt(file_path)
                        doc.extraction_method = "plaintext"
                    else:
                        raw = _extract_text_from_docx(file_path)
                        doc.extraction_method = "python-docx"

                    doc.raw_text = _clean_text(raw)
                    doc.save(update_fields=["raw_text", "extraction_method", "updated_at"])
                except ExtractionError as e:
                    doc.raw_text = ""
                    doc.extraction_method = "failed"
                    doc.save(update_fields=["raw_text", "extraction_method", "updated_at"])
                    errors.append({
                        "filename": f.name,
                        "error": str(e),
                        "error_code": e.error_code,
                    })
                    continue
                except Exception as e:
                    doc.raw_text = ""
                    doc.extraction_method = "failed"
                    doc.save(update_fields=["raw_text", "extraction_method", "updated_at"])
                    errors.append({
                        "filename": f.name,
                        "error": f"Text extraction failed: {str(e)}",
                        "error_code": "TEXT_EXTRACTION_FAILED",
                    })
                    continue

                # Create ParseRun queued and dispatch
                run = ParseRun.objects.create(
                    resume_document=doc,
                    status="queued",
                    model_name=getattr(settings, "OPENROUTER_EXTRACT_MODEL", "openai/gpt-4o-mini"),
                    prompt_version="v1",
                    temperature=float(getattr(settings, "OPENROUTER_TEMPERATURE", 0.1)),
                    requirements=requirements,  # Store requirements for async checking
                )

                if getattr(settings, "RESUME_PARSE_ASYNC", True) and not sync:
                    parse_resume_parse_run.delay(run.id, requirements=requirements)
                    results.append({
                        "filename": f.name,
                        "resume_document_id": doc.id,
                        "parse_run_id": run.id,
                        "status": "queued",
                        "requirements_check_pending": bool(requirements),  # Will check after async processing
                    })
                else:
                    # sync fallback for demos/testing
                    parse_resume_parse_run(run.id)
                    run.refresh_from_db()
                    latest_candidate = Candidate.objects.filter(parse_run=run).order_by("-created_at").first()
                    
                    # Check requirements if provided (sync mode only)
                    if requirements and latest_candidate:
                        meets, reasons = _candidate_meets_requirements(latest_candidate, requirements)
                        if not meets:
                            # Discard candidate and related data
                            discarded.append({
                                "filename": f.name,
                                "candidate_id": latest_candidate.id,
                                "reasons": reasons,
                            })
                            latest_candidate.delete()  # This will cascade delete skills, education, experience
                            results.append({
                                "filename": f.name,
                                "resume_document_id": doc.id,
                                "parse_run_id": run.id,
                                "status": run.status,
                                "discarded": True,
                                "discard_reasons": reasons,
                            })
                            continue
                    
                    results.append({
                        "filename": f.name,
                        "resume_document_id": doc.id,
                        "parse_run_id": run.id,
                        "status": run.status,
                        "candidate_id": latest_candidate.id if latest_candidate else None,
                    })

            except Exception as e:
                errors.append({
                    "filename": f.name,
                    "error": str(e),
                    "error_code": "UPLOAD_FAILED",
                })

        # Include all results with clear status indicators
        all_results = []
        for r in results:
            if r.get("discarded", False):
                r["status_type"] = "rejected"
                r["accepted"] = False
            elif r.get("duplicate", False):
                r["status_type"] = "duplicate"
                r["accepted"] = True  # Duplicates are considered accepted
            elif r.get("candidate_id"):
                r["status_type"] = "accepted"
                r["accepted"] = True
            else:
                r["status_type"] = "processed"
                r["accepted"] = True
            
            all_results.append(r)
        
        summary = {
            "total": len(validated_files),
            "successful": len(results),
            "matching": len([r for r in all_results if r.get("accepted", False) and not r.get("discarded", False)]),
            "rejected": len([r for r in all_results if r.get("discarded", False)]),
            "errors": len(errors),
            "results": all_results,  # Include all results with status indicators
            "accepted": [r for r in all_results if r.get("accepted", False) and not r.get("discarded", False)],
            "rejected": [r for r in all_results if r.get("discarded", False)],
            "errors": errors,
        }
        
        if discarded:
            summary["discarded_details"] = discarded
        
        if requirements:
            summary["requirements_applied"] = requirements
            if sync:
                summary["note"] = f"Requirements applied. {summary['matching']} candidates accepted, {summary['rejected']} rejected."
            else:
                summary["note"] = "Requirements will be checked after async processing completes. Check parse runs for final status."

        status_code = 202 if getattr(settings, "RESUME_PARSE_ASYNC", True) and not sync else 201
        return ok(summary, status=status_code)
